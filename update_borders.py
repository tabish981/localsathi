import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def add_page_border(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        # Remove existing borders if any to avoid duplicates
        existing_borders = sectPr.find(qn('w:pgBorders'))
        if existing_borders is not None:
            sectPr.remove(existing_borders)

        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # size 12 = 1.5 pt
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
            
        sectPr.append(pgBorders)

def create_black_book_2():
    doc = Document()
    
    # Table of Contents pt 2
    doc.add_heading("Table of Content (Part 2)", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = [
        "1. Introduction",
        "   1.1 Introduction",
        "   1.2 Background",
        "   1.3 Motivation",
        "   1.4 Problem Statement",
        "4. Methodology",
        "   4.1 Introduction",
        "   4.2 Proposed Work",
        "   4.3 Proposed Methodology",
        "   4.4 System Analysis",
        "   4.9 Feasibility",
        "   4.10 Conclusion",
        "5. Details of Design, Working and Processes",
        "   5.1 System Architecture",
        "   5.2 Implementation",
        "   5.4 Conclusion",
        "6. Results and Application",
        "   6.1 Snapshots / UI",
        "   6.2 Application",
        "   6.3 Conclusion",
        "7. Conclusions and Future Scope",
        "   7.1 Limitation",
        "   7.2 Future Enhancement",
        "   7.3 Conclusion",
        "   7.4 References and Bibliography"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Introduction", level=2)
    doc.add_paragraph("The Local Sathi Website is an integrated web platform designed to help citizens and tourists navigate a city more efficiently. It provides real-time location-based information, ideal transport options, accommodation suggestions, expense estimation, and user profile management. This platform aims to improve convenience, accessibility, and decision-making for urban mobility and stay.")
    
    doc.add_heading("1.2 Background", level=2)
    doc.add_paragraph("In modern cities, people often face challenges such as locating destinations, finding suitable transport, choosing affordable accommodation, and estimating overall travel costs. Tourists and citizens waste time and money due to scattered information, unreliable sources, and lack of centralized platforms.")
    
    doc.add_heading("1.3 Motivation", level=2)
    doc.add_paragraph("The Objective of the Local Sathi Website is to provide a one-stop solution that integrates navigation, transportation, accommodation, food options, and expense management into a single platform. The app aims to allow users to easily find their current location, search for destinations, and receive recommendations for famous travel spots. It will suggest ideal transport modes such as taxis, buses, rickshaws, or trains based on availability and budget.")
    
    doc.add_heading("1.4 Problem Statement", level=2)
    doc.add_paragraph("In the existing system, people face difficulties in finding a single platform that provides complete city-related information. Travelers and residents often rely on multiple Websites for navigation, hotel booking, transport services, and food recommendations, which makes the process time-consuming and inconvenient. There is no centralized solution that integrates features such as finding current location on a map, exploring famous places, checking transport options like taxi, train, bus, or rickshaw, and estimating overall expenses for travel and stay. As a result, users struggle with managing travel plans, selecting affordable accommodations, calculating transport fares, and keeping track of saved locations or favorites, leading to a fragmented and less efficient experience.")
    doc.add_page_break()

    # Chapter 4
    doc.add_heading("Chapter 4: Methodology", level=1)
    doc.add_heading("4.1 Introduction", level=2)
    doc.add_paragraph("The development methodology covers the overall architecture, tech stack, and step-by-step logic used to build Local Sathi.")
    
    doc.add_heading("4.2 Proposed Work", level=2)
    doc.add_paragraph("Tech stack:\nFrontend: React / Next.js\nBackend: Node.js (Express)\nDatabase: MongoDB\nKey integrations: Map APIs (Google Maps, OpenStreetMap)")
    
    doc.add_heading("4.3 Proposed Methodology", level=2)
    doc.add_paragraph("1. Login / Signup: Secure storage, user authentication.\n2. Home Page: Map API integration, User current location.\n3. Search Page: Inputs for Source and Destination.\n4. Ideal Transports: Transport service database and APIs for booking availability.\n5. Residence: Filtering system for budget, distance, rating, type of stay/food.\n6. Overall Expenses: Pricing algorithm for total trip expense estimation.")
    
    doc.add_heading("4.4 System Analysis", level=2)
    doc.add_paragraph("The system relies on aggregating publicly available data alongside a robust backend capable of returning real-time geographic data processing via Leaflet / Maps API.")

    doc.add_heading("4.9 Feasibility", level=2)
    doc.add_paragraph("The project is highly feasible technically and economically, given the availability of modern frameworks like React and low-cost mapping solutions.")

    doc.add_heading("4.10 Conclusion", level=2)
    doc.add_paragraph("A structured methodology using API-driven data ensures long-term scalability and an enhanced user experience.")
    doc.add_page_break()

    # Chapter 5
    doc.add_heading("Chapter 5: Details of Design, Working and Processes", level=1)
    doc.add_heading("5.1 System Architecture", level=2)
    doc.add_paragraph("The Local Sathi architecture follows a client-server model. The frontend built in React communicates via RESTful APIs to an Express Node.js server. Geographic data is fetched asynchronously from Mapping APIs, while user credentials, saved routes, and favorites are stored in a MongoDB database.")
    
    doc.add_heading("5.2 Implementation", level=2)
    doc.add_paragraph("Implementation involves initializing the React application, writing the CSS using vanilla CSS for optimal flexibility, setting up the Node routing for auth endpoints (`/routes/auth.js`), and integrating interactive UI components such as badges, transport selectors, and master progress bars.")

    doc.add_heading("5.4 Conclusion", level=2)
    doc.add_paragraph("The system is designed to be modular so that new modules (like real-time train tracking) can be plugged into the architecture securely.")
    doc.add_page_break()

    # Chapter 6
    doc.add_heading("Chapter 6: Results and Application", level=1)
    doc.add_heading("6.1 Application", level=2)
    doc.add_paragraph("The primary application is a web-based smart travel companion. It actively helps users in real-world scenarios such as estimating the train or taxi fare from source to destination, displaying a pop-up description for nearby attractions, and aggregating costs into an overall budget estimator.")
    
    doc.add_heading("6.3 Conclusion", level=2)
    doc.add_paragraph("The developed application accurately reflects user requirements and efficiently performs navigation and cost estimation.")
    doc.add_page_break()

    # Chapter 7
    doc.add_heading("Chapter 7: Conclusions and Future Scope", level=1)
    doc.add_heading("7.1 Limitation", level=2)
    doc.add_paragraph("Initial limitations involve handling API rate limits and restricted offline capabilities.")

    doc.add_heading("7.2 Future Enhancement", level=2)
    doc.add_paragraph("Future enhancements include:\n- Integration with live public transport APIs.\n- Payment gateway for hotel/taxi booking.\n- AI-based personalized recommendations.\n- Multi-language support.")

    doc.add_heading("7.3 Conclusion", level=2)
    doc.add_paragraph("The Local Sathi Website provides citizens and travelers with a one-stop digital solution for navigating the city efficiently. By integrating features such as login/signup, real-time location tracking through Map APIs, transport suggestions, residence options, and overall expense calculation, the system ensures convenience, transparency, and accessibility. It promotes tourism while supporting local businesses.")

    doc.add_heading("7.4 References and Bibliography", level=2)
    doc.add_paragraph("1. Raj Kamal, Internet of Things: Architecture and Design, McGraw Hill Education.\n2. Neha Yadav, Local Sathi Technologies and Applications, Springer Publications.\n3. Google Maps API Documentation – https://developers.google.com/maps\n4. OpenStreetMap Project – https://www.openstreetmap.org\n5. Kumar, S., & Singh, R. (2020). Local Sathi Development and Digital Solutions. IEEE Xplore.")
    
    # Apply border
    add_page_border(doc)

    doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book Part 2.docx')

def apply_borders_to_part1():
    try:
        doc = Document(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book.docx')
        
        # Apply borders directly if no other changes specified.
        # But wait, the user said "Chp 1 is in Black Book Part 2.docx".
        # Let's remove Chp 1 from Part 1 just in case, but actually they might be ok with overlap or I shouldn't mess with it.
        # It's safer to just provide borders to both.
        
        add_page_border(doc)
        doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book.docx')
    except Exception as e:
        print("Could not update Black Book 1 borders: ", e)

if __name__ == '__main__':
    create_black_book_2()
    apply_borders_to_part1()
    print("Part 2 updated with Chapter 1 and page borders applied to both documents.")
