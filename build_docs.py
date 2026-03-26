import os
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def add_page_border(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        # Remove existing borders if any
        existing_borders = sectPr.find(qn('w:pgBorders'))
        if existing_borders is not None:
            sectPr.remove(existing_borders)

        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '12')  # size 12 = 1.5 pt
            border.set(qn('w:space'), '24')
            border.set(qn('w:color'), 'auto')
            pgBorders.append(border)
            
        sectPr.append(pgBorders)

def create_black_book_1():
    doc = Document()
    
    # Title Page
    title = doc.add_paragraph("Anjuman-I-Islam’s\nM. H. Saboo Siddik Polytechnic\n8, M.H.Saboo Siddik Polytechnic Road, Mumbai 400008")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("PROJECT REPORT ON", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading("Local Sathi (A Digital Solution for Smart Cities)", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("BY\nAnsari Tauheed (23150360424)\nTabish Ansari (23150360426)\nAtaur Rahman (23150360428)\nBasit Shaikh (23150360471)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("UNDER THE GUIDANCE OF\nMs. Uzma Alam Khan / Ms. Farhanaaz Sayed").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Maharashtra State Board of Technical Education (MS-BTE)\nMumbai (Autonomous)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Certificate
    doc.add_heading("CERTIFICATE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("This is to certify that Mr. Ansari Tauheed, Tabish Ansari, Ataur Rahman, Basit Shaikh of 5TH Semester of Diploma in Computer Engineering of Institute M.H Saboo Siddik Polytechnic has successfully completed project work in subject for the academic year 2025-26 as prescribed in the K-Scheme Curriculum.")
    doc.add_paragraph("HEAD OF DEPARTMENT:                     INTERNAL EXAMINER:\nDATE:                                   DATE:\n\nSIGN OF THE GUIDE:                      EXTERNAL EXAMINER:\nDATE:                                   DATE:\n\nPRINCIPAL:\nDATE:")
    doc.add_page_break()

    # Acknowledgement
    doc.add_heading("Acknowledgement", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("We would like to express our sincere gratitude to our mentor Ms. Farhanaaz Sayed / Ms. Uzma Alam Khan for her valuable inputs, timely suggestions, and encouragement, which helped us stay motivated and focused during our Project.")
    doc.add_paragraph("We are equally grateful to Ms. Bushra Shaikh, Head of Department, Computer Engineering (Unaided), M. H. Saboo Siddik Polytechnic, for her continuous support and encouragement. We extend our heartful gratitude to Dr. A. K. Kureshi, Principal, M. H. Saboo Siddik Polytechnic, for providing the necessary support and infrastructure to successfully complete this Project.")
    doc.add_page_break()

    # Abstract
    doc.add_heading("Abstract", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("The Local Sathi Guide is a comprehensive Website designed to simplify urban exploration and resource management, offering a seamless experience from initial login/signup to detailed travel planning. The Home Page utilizes real-time map integration to establish the user's current location and provides quick access to essential services via specialized categories like Residences and Ideal Transports. Users can leverage the Search Page to plan routes between locations, discover recommended famous destinations, and view integrated options for local transport (taxis, trains, buses) and diverse lodging options (hotels, motels, dhabas). Crucially, the Website provides an Overall Expenses overview, summarizing costs like accommodation rent and transport fares for transparent budget planning, all managed through a dedicated Profile where settings and favorite locations are saved.")
    doc.add_page_break()

    # Chapter 2
    doc.add_heading("Chapter 2: Literature Survey", level=1)
    doc.add_heading("2.1 Introduction", level=2)
    doc.add_paragraph("The growth of smart cities has prompted extensive research into smart transportation, IoT integrations, and web portals that streamline tourism and accommodation.")
    
    doc.add_heading("2.2 Research Papers", level=2)
    doc.add_paragraph("1. A Global and Dynamic Route Planning Application for Smart Transportation (Ning Sun et al., 2020) - Discusses optimization of global path planning based on traffic.")
    doc.add_paragraph("2. IoT and Microservice Architecture for Multimobility in a Smart City (Cristian Lai) - Microservices for scalable multi-mobility systems.")
    doc.add_paragraph("3. Gamification of Citizen Participation (Caroline L.) - Explores enhancing user experiences through dynamic participatory tools.")
    doc.add_paragraph("4. Comprehensive Review on Development of Smart Cities Using Industry 4.0 Technologies - Cloud processing and data integration.")
    
    doc.add_heading("2.3 References", level=2)
    doc.add_paragraph("See final chapter for full bibliography of IEEE journals and online API documentations.")
    
    doc.add_heading("2.4 Conclusion", level=2)
    doc.add_paragraph("Literature shows a strong need for integrated platforms that do not just map a route but predict cost, consolidate booking services, and offer dynamic web applications using React and map APIs.")
    doc.add_page_break()

    # Chapter 3
    doc.add_heading("Chapter 3: Scope of the Project", level=1)
    doc.add_heading("3.1 Introduction", level=2)
    doc.add_paragraph("This chapter defines the boundary of the Local Sathi project, including its capabilities, focus areas, and potential limitations.")
    
    doc.add_heading("3.2 Scope", level=2)
    doc.add_paragraph("The Scope of the project covers user account management, interactive map integration for navigation, transport suggestions with estimated fares, and listings of accommodation and food services. The app will also include an expense calculator to help users budget their trips effectively. It is designed to assist both residents and tourists by recommending popular places and nearby services.")
    
    doc.add_heading("3.3 Objective", level=2)
    doc.add_paragraph("- Provide an interactive map to determine the user’s current location.\n- Suggest transportation options with cost and time estimates.\n- Recommend nearby hotels, food stalls, lounges, and motels.\n- Calculate estimated expenses (transport + accommodation).\n- Allow users to save locations and preferences in their profiles.")
    
    doc.add_heading("3.4 Advantages", level=2)
    doc.add_paragraph("1. Centralized Information Platform\n2. Real-Time Location Tracking\n3. Improved Decision-Making\n4. Tourism Boost\n5. User-Friendly Interface\n6. Cost Estimation")
    
    doc.add_heading("3.5 Disadvantages", level=2)
    doc.add_paragraph("1. Dependency on Internet & APIs\n2. Initial Limited Data\n3. Maintenance Cost\n4. Privacy & Security Risks\n5. Competition with Existing Apps")
    
    doc.add_heading("3.6 Conclusion", level=2)
    doc.add_paragraph("While there are a few drawbacks concerning API dependency and maintenance, the advantages heavily outweigh them. By bringing together transportation, living, and finance tools into a single platform, Local Sathi is poised to provide high value to the smart city ecosystem.")

    add_page_border(doc)
    
    try:
        doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book.docx')
    except PermissionError:
        print("Warning: Black Book.docx is open. Saving as Black Book_Updated.docx")
        doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book_Updated.docx')


def create_black_book_2():
    doc = Document()
    
    # Table of Contents
    doc.add_heading("Table of Content", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = [
        "1. Introduction", 
        "   1.1 Introduction",
        "   1.2 Background",
        "   1.3 Motivation",
        "   1.4 Problem Statement",
        "2. Literature Survey",
        "   2.1 Introduction",
        "   2.2 Research Papers",
        "   2.3 References",
        "   2.4 Conclusion",
        "3. Scope of the Project",
        "   3.1 Introduction",
        "   3.2 Scope",
        "   3.3 Objective",
        "   3.4 Advantages",
        "   3.5 Disadvantages",
        "   3.6 Conclusion",
        "4. Methodology",
        "   4.1 Introduction",
        "   4.2 Proposed Work",
        "   4.3 Proposed Methodology",
        "   4.4 System Analysis",
        "   4.9 Feasibility",
        "   4.10 Conclusion",
        "5. Details of Design, Working and Processes",
        "   5.1 System Architecture",
        "   5.2 Implementation",
        "   5.4 Conclusion",
        "6. Results and Application",
        "   6.1 Snapshots / UI",
        "   6.2 Application",
        "   6.3 Conclusion",
        "7. Conclusions and Future Scope",
        "   7.1 Limitation",
        "   7.2 Future Enhancement",
        "   7.3 Conclusion",
        "   7.4 References and Bibliography"
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # Chapter 1
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Introduction", level=2)
    doc.add_paragraph("The Local Sathi Website is an integrated web platform designed to help citizens and tourists navigate a city more efficiently. It provides real-time location-based information, ideal transport options, accommodation suggestions, expense estimation, and user profile management. This platform aims to improve convenience, accessibility, and decision-making for urban mobility and stay.")
    
    doc.add_heading("1.2 Background", level=2)
    doc.add_paragraph("In modern cities, people often face challenges such as locating destinations, finding suitable transport, choosing affordable accommodation, and estimating overall travel costs. Tourists and citizens waste time and money due to scattered information, unreliable sources, and lack of centralized platforms.")
    
    doc.add_heading("1.3 Motivation", level=2)
    doc.add_paragraph("The Objective of the Local Sathi Website is to provide a one-stop solution that integrates navigation, transportation, accommodation, food options, and expense management into a single platform. The app aims to allow users to easily find their current location, search for destinations, and receive recommendations for famous travel spots. It will suggest ideal transport modes such as taxis, buses, rickshaws, or trains based on availability and budget.")
    
    doc.add_heading("1.4 Problem Statement", level=2)
    doc.add_paragraph("In the existing system, people face difficulties in finding a single platform that provides complete city-related information. Travelers and residents often rely on multiple Websites for navigation, hotel booking, transport services, and food recommendations, which makes the process time-consuming and inconvenient. There is no centralized solution that integrates features such as finding current location on a map, exploring famous places, checking transport options like taxi, train, bus, or rickshaw, and estimating overall expenses for travel and stay. As a result, users struggle with managing travel plans, selecting affordable accommodations, calculating transport fares, and keeping track of saved locations or favorites, leading to a fragmented and less efficient experience.")
    doc.add_page_break()

    # Chapter 4
    doc.add_heading("Chapter 4: Methodology", level=1)
    doc.add_heading("4.1 Introduction", level=2)
    doc.add_paragraph("The development methodology covers the overall architecture, tech stack, and step-by-step logic used to build Local Sathi.")
    
    doc.add_heading("4.2 Proposed Work", level=2)
    doc.add_paragraph("Tech stack:\nFrontend: React / Next.js\nBackend: Node.js (Express)\nDatabase: MongoDB\nKey integrations: Map APIs (Google Maps, OpenStreetMap)")
    
    doc.add_heading("4.3 Proposed Methodology", level=2)
    doc.add_paragraph("1. Login / Signup: Secure storage, user authentication.\n2. Home Page: Map API integration, User current location.\n3. Search Page: Inputs for Source and Destination.\n4. Ideal Transports: Transport service database and APIs for booking availability.\n5. Residence: Filtering system for budget, distance, rating, type of stay/food.\n6. Overall Expenses: Pricing algorithm for total trip expense estimation.")
    
    doc.add_heading("4.4 System Analysis", level=2)
    doc.add_paragraph("The system relies on aggregating publicly available data alongside a robust backend capable of returning real-time geographic data processing via Leaflet / Maps API.")

    doc.add_heading("4.9 Feasibility", level=2)
    doc.add_paragraph("The project is highly feasible technically and economically, given the availability of modern frameworks like React and low-cost mapping solutions.")

    doc.add_heading("4.10 Conclusion", level=2)
    doc.add_paragraph("A structured methodology using API-driven data ensures long-term scalability and an enhanced user experience.")
    doc.add_page_break()

    # Chapter 5
    doc.add_heading("Chapter 5: Details of Design, Working and Processes", level=1)
    doc.add_heading("5.1 System Architecture", level=2)
    doc.add_paragraph("The Local Sathi architecture follows a client-server model. The frontend built in React communicates via RESTful APIs to an Express Node.js server. Geographic data is fetched asynchronously from Mapping APIs, while user credentials, saved routes, and favorites are stored in a MongoDB database.")
    
    doc.add_heading("5.2 Implementation", level=2)
    doc.add_paragraph("Implementation involves initializing the React application, writing the CSS using vanilla CSS for optimal flexibility, setting up the Node routing for auth endpoints (`/routes/auth.js`), and integrating interactive UI components such as badges, transport selectors, and master progress bars.")

    doc.add_heading("5.4 Conclusion", level=2)
    doc.add_paragraph("The system is designed to be modular so that new modules (like real-time train tracking) can be plugged into the architecture securely.")
    doc.add_page_break()

    # Chapter 6
    doc.add_heading("Chapter 6: Results and Application", level=1)
    doc.add_heading("6.1 Application", level=2)
    doc.add_paragraph("The primary application is a web-based smart travel companion. It actively helps users in real-world scenarios such as estimating the train or taxi fare from source to destination, displaying a pop-up description for nearby attractions, and aggregating costs into an overall budget estimator.")
    
    doc.add_heading("6.3 Conclusion", level=2)
    doc.add_paragraph("The developed application accurately reflects user requirements and efficiently performs navigation and cost estimation.")
    doc.add_page_break()

    # Chapter 7
    doc.add_heading("Chapter 7: Conclusions and Future Scope", level=1)
    doc.add_heading("7.1 Limitation", level=2)
    doc.add_paragraph("Initial limitations involve handling API rate limits and restricted offline capabilities.")

    doc.add_heading("7.2 Future Enhancement", level=2)
    doc.add_paragraph("Future enhancements include:\n- Integration with live public transport APIs.\n- Payment gateway for hotel/taxi booking.\n- AI-based personalized recommendations.\n- Multi-language support.")

    doc.add_heading("7.3 Conclusion", level=2)
    doc.add_paragraph("The Local Sathi Website provides citizens and travelers with a one-stop digital solution for navigating the city efficiently. By integrating features such as login/signup, real-time location tracking through Map APIs, transport suggestions, residence options, and overall expense calculation, the system ensures convenience, transparency, and accessibility. It promotes tourism while supporting local businesses.")

    doc.add_heading("7.4 References and Bibliography", level=2)
    doc.add_paragraph("1. Raj Kamal, Internet of Things: Architecture and Design, McGraw Hill Education.\n2. Neha Yadav, Local Sathi Technologies and Applications, Springer Publications.\n3. Google Maps API Documentation – https://developers.google.com/maps\n4. OpenStreetMap Project – https://www.openstreetmap.org\n5. Kumar, S., & Singh, R. (2020). Local Sathi Development and Digital Solutions. IEEE Xplore.")
    
    add_page_border(doc)
    
    try:
        doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book Part 2.docx')
    except PermissionError:
        print("Warning: Black Book Part 2.docx is open. Saving as Black Book Part 2_Updated.docx")
        doc.save(r'c:\Users\tabish Ansari\OneDrive\Desktop\localsathi\Black Book Part 2_Updated.docx')

if __name__ == '__main__':
    create_black_book_1()
    create_black_book_2()
    print("Documents successfully restructured and page borders added.")
